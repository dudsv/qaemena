import os
import re
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
from tkinter import Tk, filedialog, simpledialog, messagebox, Button, Label
from tkinter import ttk
from pathlib import Path
from playwright.sync_api import sync_playwright

# ---------------------- AUXILIARES ----------------------

def clean_text(text):
    if not isinstance(text, str):
        return ""
    # Remove links (http, www, etc.) e parênteses vazios
    text = re.sub(
        r'\s*(?:\([^()]*?(?:https?://|www\.|/)[^()]*?\)|(?:https?://|www\.|/)[^\s()]+)\s*',
        ' ',
        text
    )
    text = re.sub(r'\(\s*\)', '', text)
    # Colapsa espaços extras, mas preserva espaço antes de pontuação
    text = re.sub(r'\s+|(\s+)(?=[.,!?:;])', lambda m: '' if m.group(1) else ' ', text)
    return text.strip()


def extrair_metadados(soup):
    return {
        "Title Tag": soup.title.string.strip() if soup.title else "",
        "Meta Description": soup.find("meta", {"name": "description"})['content'].strip()
                             if soup.find("meta", {"name": "description"}) else "",
        "Open Graph Title": soup.find("meta", {"property": "og:title"})['content'].strip()
                             if soup.find("meta", {"property": "og:title"}) else "",
        "Open Graph Description": soup.find("meta", {"property": "og:description"})['content'].strip()
                                  if soup.find("meta", {"property": "og:description"}) else ""
    }


def extrair_alt_tags(soup):
    return [img.get("alt", "").strip() for img in soup.find_all("img") if img.get("alt")]


def coletar_elementos_html(main):
    # Remove footer
    footer = main.find('footer')
    if footer:
        footer.decompose()

    elementos = []
    # Headings, bold e italic (mantém seu filtro original)
    for i in range(1, 7):
        for tag in main.find_all(f'h{i}'):
            texto = clean_text(tag.get_text(" ", strip=True))
            if texto:
                elementos.append(['Heading', f'h{i}', texto, ''])
    for tag in main.find_all(['strong', 'b']):
        texto = clean_text(tag.get_text(" ", strip=True))
        if texto:
            elementos.append(['Bold', '', texto, ''])
    for tag in main.find_all(['em', 'i']):
        texto = clean_text(tag.get_text(" ", strip=True))
        if texto:
            elementos.append(['Italic', '', texto, ''])

    # Hyperlinks
    ignorar = {
        '/#copy_link', '/#x', '/#facebook', '/#whatsapp',
        'https://www.addtoany.com',
        'https://www.nestle.be/fr/info/yourdata',
        'https://www.onetrust.com/products/cookie-consent/'
    }
    for tag in main.find_all('a', href=True):
        href = tag['href'].strip()
        texto = clean_text(tag.get_text(" ", strip=True))
        # filtra anchors vazios ou indesejados
        if (not href
            or href.startswith('#')
            or href in ignorar
            or texto.lower() in ('previous next', 'anterior siguiente')):
            continue
        elementos.append([
            'Hyperlink',  # Coluna A
            '',           # Coluna B em branco
            texto,        # Coluna C
            href          # Coluna D
        ])

    return pd.DataFrame(
        elementos,
        columns=['Definição', 'Heading', 'Texto', 'Link']
    )


def carregar_texto_docx(path):
    doc = Document(path)
    return [clean_text(p.text) for p in doc.paragraphs if p.text.strip()]


def carregar_texto_url(url):
    """
    Carrega a página via Playwright, aguarda networkidle, faz scroll,
    extrai o título do accordion “Puntuación Veterinaria” pela <a class="accordion--text-v2">,
    depois extrai a tabela dentro da div específica e todos os blocos de texto úteis,
    ignorando apenas o header “Previous Next” que aparece junto ao título.
    """
    # 1) Render JS e aguardar carregamento
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(url, timeout=60000)
        page.wait_for_load_state("networkidle")
        page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
        page.wait_for_timeout(1000)
        html = page.content()
        browser.close()

    soup = BeautifulSoup(html, 'html.parser')
    textos = []

    # 2) Captura título do accordion via <a class="accordion--text-v2">
    for a_tag in soup.find_all("a", class_="accordion--text-v2"):
        raw = clean_text(a_tag.get_text(" ", strip=True))
        if raw:
            textos.append(raw)

    # 3) Localiza a div do accordion “Puntuación Veterinaria” e extrai a tabela, se existir
    accordion_div = soup.find(
        "div",
        {
            "class": "text-image--text-wrapper col-12 col-xl-5 order-3 order-xl-2"
        }
    )
    if accordion_div:
        tabela = accordion_div.find("table", {"class": "breed-table breed-table-col-2"})
        if tabela:
            for linha in tabela.find_all("tr"):
                células = linha.find_all("td")
                if len(células) >= 2:
                    chave = clean_text(células[0].get_text(" ", strip=True))
                    valor = clean_text(células[1].get_text(" ", strip=True))
                    if chave and valor:
                        textos.append(f"{chave}: {valor}")

    # 4) Extrai metadados e define main
    metadados = extrair_metadados(soup)
    main = soup.find("main")
    if not main or len(main.get_text(strip=True)) < 50:
        main = soup.body

    # Remove footer, nav, scripts, estilos e noscript do main
    for tag in main.find_all(['footer', 'nav', 'script', 'style', 'noscript', 'menu', 'dialog']):
        tag.decompose()

    # Extrai lista de alt tags para comparação
    alt_tags = extrair_alt_tags(main)

    # Extrai imagens para nova aba: URL e Texto Alt
    imagens = []
    for img in main.find_all("img"):
        alt = img.get("alt", "").strip()
        src = img.get("src", "").strip()
        # incluir só se vier do style que você quer
        if alt and src and "/styles/ttt_image_510/" in src:
            imagens.append((src, clean_text(alt)))


    # 5) Extrai h1–h6, p, li, span e div do main, mas ignora blocos que contenham "Previous Next"
    blocos = main.find_all(['h1','h2','h3','h4','h5','h6','p','li','span','div'])
    for tag in blocos:
        raw_text = tag.get_text(" ", strip=True)
        if "Previous Next" in raw_text or "Anterior Siguiente" in raw_text:
            continue
        txt = clean_text(raw_text)
        if txt:
            textos.append(txt)

    titulo = soup.title.string.strip() if soup.title else "pagina"
    return textos, main, metadados, alt_tags, titulo, imagens


def safe_best_match(query, candidates):
    """
    Retorna (melhor_texto, similaridade) ou ("", 0.0).
    Evita erro de "empty vocabulary" quando não há texto útil.
    """
    query = clean_text(query.lower())
    candidatos_limpos = [clean_text(c.lower()) for c in candidates if clean_text(c)]
    if not query or not candidatos_limpos:
        return "", 0.0

    try:
        vectorizer = TfidfVectorizer(stop_words='english')
        corpus = [query] + candidatos_limpos
        tfidf = vectorizer.fit_transform(corpus)
        sims = cosine_similarity(tfidf[0:1], tfidf[1:]).flatten()
        idx_max = sims.argmax()
        return candidates[idx_max], float(sims[idx_max])
    except ValueError:
        return "", 0.0


def comparar_textos(lista_docx, lista_html, metadados, alt_tags):
    resultados = []

    for texto_doc in lista_docx:
        if not texto_doc.strip():
            continue
        texto_limpo = clean_text(texto_doc.strip().lower())
        prefixos_ignorados = [
            "in dit artikel", "title tag:", "meta description:", "og title:", "og description:",
            "[alt text da imagem]", "alt tag :", "title tag", "meta description",
            "open graph title", "open graph description", "-- meta --", "en:", "be-fr:",
            "guide des races de chiens", "alt-tag:", "-- meta –", "title tag", "Etiqueta alt: ", "Etiqueta alt:"
        ]
        if any(texto_limpo.startswith(p) for p in prefixos_ignorados):
            continue

        # Bloco "alt-tag"
        if texto_limpo.startswith("alt-tag"):
            original_alt = texto_doc.split(":", 1)[1].strip()
            match_text, score = safe_best_match(original_alt, alt_tags)
            if score >= 0.85:
                status = "Exact"
            elif score >= 0.75:
                status = "Similar"
            elif score >= 0.4:
                status = "Partial"
            else:
                status = "Missing"
            resultados.append({
                "Document Text": original_alt,
                "Webpage Match": match_text,
                "Status": status,
                "Similarity": round(score * 100, 1)
            })
            continue

        # Checa match exato em metadados
        tipo_meta = next((k for k, v in metadados.items() if v and texto_doc.strip() == v.strip()), None)
        if not tipo_meta:
            melhor_meta = ""
            melhor_score_mt = 0.0
            for k, v in metadados.items():
                if not v.strip():
                    continue
                _, sim = safe_best_match(texto_doc, [v])
                if sim > melhor_score_mt:
                    melhor_score_mt = sim
                    melhor_meta = k
            if melhor_score_mt > 0.85:
                tipo_meta = melhor_meta

        if tipo_meta:
            sim_meta = safe_best_match(texto_doc, [metadados[tipo_meta]])[1]
            if sim_meta >= 0.85:
                status = "Exact"
            elif sim_meta >= 0.75:
                status = "Similar"
            elif sim_meta >= 0.4:
                status = "Partial"
            else:
                status = "Missing"
            resultados.append({
                "Document Text": texto_doc,
                "Webpage Match": metadados[tipo_meta],
                "Status": status,
                "Similarity": round(sim_meta * 100, 1)
            })
            continue

        # Senão, compara contra blocos HTML
        match_html, score_html = "", 0.0
        if lista_html:
            match_html, score_html = safe_best_match(texto_doc, lista_html)
        if score_html >= 0.85:
            status = "Exact"
        elif score_html >= 0.75:
            status = "Similar"
        elif score_html >= 0.4:
            status = "Partial"
        else:
            status = "Missing"
        resultados.append({
            "Document Text": texto_doc,
            "Webpage Match": match_html,
            "Status": status,
            "Similarity": round(score_html * 100, 1)
        })

    return pd.DataFrame(resultados)


def gerar_resumo(df):
    total = len(df)
    resumo = df["Status"].value_counts().reindex(
        ["Exact", "Similar", "Partial", "Missing"], fill_value=0
    )
    porcentagens = (resumo / total * 100).round(1) if total > 0 else [0, 0, 0, 0]
    df_resumo = pd.DataFrame({
        "Status": resumo.index,
        "Quantidade": resumo.values,
        "Porcentagem": porcentagens.values
    })
    df_resumo.loc[len(df_resumo.index)] = ["TOTAL", total, f"{100 if total > 0 else 0}%"]
    return df_resumo


# ---------------------- EXTRAÇÃO DE TEXTO EM NEGRITO ----------------------
def extract_bold_phrases(doc_path):
    """Agrupa runs, tabelas e caixas de texto para capturar frases completas em negrito."""
    doc = Document(doc_path)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    def runs_to_phrases(runs):
        texts, curr = [], []
        for r in runs:
            rPr = r.find("w:rPr", ns)
            bold = rPr is not None and rPr.find("w:b", ns) is not None
            t = r.find("w:t", ns)
            if bold and t is not None and t.text.strip():
                curr.append(t.text.strip())
            else:
                if curr:
                    texts.append(" ".join(curr))
                    curr = []
        if curr:
            texts.append(" ".join(curr))
        return texts

    phrases = []
    for para in doc.paragraphs:
        phrases += runs_to_phrases([r._element for r in para.runs])
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    phrases += runs_to_phrases([r._element for r in para.runs])
    for txbx in doc.element.body.findall('.//w:txbxContent', ns):
        for p in txbx.findall('.//w:p', ns):
            phrases += runs_to_phrases(p.findall('.//w:r', ns))

    seen, unique = set(), []
    for ph in phrases:
        norm = re.sub(r"\s+", " ", ph).strip()
        if norm and norm not in seen:
            seen.add(norm)
            unique.append(norm)
    return unique
# ---------------------------------------------------------------------------


def salvar_em_excel(df_comparacao, df_resumo, df_elementos, df_imagens, metadados, page_url, nome_arquivo="comparacao_resultado.xlsx", word_path=None):
    wb = Workbook()

    def estilizar(ws, colorir=False):
        # Estiliza cabeçalho
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        # Ajusta largura das colunas
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = 25
        # Pinta as células conforme status, se colorir=True
        if colorir:
            cores = {
                "Exact": "C6EFCE",
                "Similar": "FFEB9C",
                "Partial": "F4B084",
                "Missing": "F8CBAD"
            }
            status_idx = next((i + 1 for i, cell in enumerate(ws[1]) if cell.value == "Status"), None)
            if status_idx:
                for row in ws.iter_rows(min_row=2):
                    status = row[status_idx - 1].value
                    cor = cores.get(status, "FFFFFF")
                    for cell in row:
                        cell.fill = PatternFill(start_color=cor, end_color=cor, fill_type="solid")

    # Aba 1: Comparacao
    aba1 = wb.active
    aba1.title = "Comparacao"
    for row in dataframe_to_rows(df_comparacao, index=False, header=True):
        aba1.append(row)
    estilizar(aba1, colorir=True)

    # Aba 2: Resumo
    aba2 = wb.create_sheet("Resumo")
    for row in dataframe_to_rows(df_resumo, index=False, header=True):
        aba2.append(row)
    estilizar(aba2)

    # Aba 3: Elementos da Pagina
    aba3 = wb.create_sheet("Elementos da Pagina")
    for row in dataframe_to_rows(df_elementos, index=False, header=True):
        aba3.append(row)
    estilizar(aba3)

    # Aba 4: Imagens
    aba4 = wb.create_sheet("Imagens")
    for row in dataframe_to_rows(df_imagens, index=False, header=True):
        aba4.append(row)
    estilizar(aba4)

 # Aba 5: Metadados
    aba5 = wb.create_sheet("Metadados")
    headers_meta = ["Page URL", "Title Tag", "Meta Description", "OG Title", "OG Description"]
    aba5.append(headers_meta)
    for cell in aba5[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    aba5.append([
        page_url,
        metadados.get("Title Tag", ""),
        metadados.get("Meta Description", ""),
        metadados.get("Open Graph Title", ""),
        metadados.get("Open Graph Description", "")
    ])
    for col, width in zip(["A","B","C","D","E"], [50,30,50,30,50]):
        aba5.column_dimensions[col].width = width

    # ---- Nova aba: Italic-Bold Check ----
    if word_path:
        bold_list = extract_bold_phrases(word_path)
        els = df_elementos["Texto"].astype(str).tolist()
        linhas = [
            {
                "From doc": t,
                "At URL": t if t in els else "",
                "Status": "Present" if t in els else "Missing",
            }
            for t in bold_list
        ]
        if linhas:
            aba6 = wb.create_sheet("Italic-Bold Check")
            df_bold = pd.DataFrame(linhas)
            for row in dataframe_to_rows(df_bold, index=False, header=True):
                aba6.append(row)
            for cell in aba6[1]:
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            for col in aba6.columns:
                aba6.column_dimensions[col[0].column_letter].width = 40
            cores_bold = {"Present": "C6EFCE", "Missing": "F8CBAD"}
            status_idx = 3
            for row in aba6.iter_rows(min_row=2):
                status = row[status_idx - 1].value
                cor = cores_bold.get(status, "FFFFFF")
                for cell in row:
                    cell.fill = PatternFill(start_color=cor, end_color=cor, fill_type="solid")

    wb.save(nome_arquivo)

# ---------------------- INTERFACE ----------------------

def executar_comparador():
    root = Tk()
    root.title("Comparador de Documentos")
    root.geometry("450x350")

    Label(root, text="Escolha o modo de comparação:", font=("Arial", 12)).pack(pady=10)
    Button(root, text="Single Page", width=30, command=lambda: comparar_um(root)).pack(pady=10)
    Button(root, text="Multiple Pages", width=30, command=lambda: comparar_varios(root)).pack(pady=10)
    root.mainloop()


def comparar_um(root):
    docx_path = filedialog.askopenfilename(
        title="Selecione o arquivo .docx", filetypes=[("Word", "*.docx")]
    )
    if not docx_path:
        return
    url = simpledialog.askstring("URL", "Insira a URL correspondente:")
    if not url:
        return
    pasta = filedialog.askdirectory(title="Selecione a pasta de saída")
    if not pasta:
        return

    try:
        docx_txt = carregar_texto_docx(docx_path)
        html_txt, main, meta, alts, titulo, imagens = carregar_texto_url(url)
        df1 = comparar_textos(docx_txt, html_txt, meta, alts)
        df1 = df1[df1["Document Text"].str.strip() != ""]
        df2 = gerar_resumo(df1)
        df3 = coletar_elementos_html(main)
        df4 = pd.DataFrame(imagens, columns=["Image URL", "Image Alt"])
        nome = re.sub(r'[\\/:*?"<>|]', '', titulo)[:50]
        salvar_em_excel(
            df1,
            df2,
            df3,
            df4,
            meta,
            url,
            os.path.join(pasta, f"comparacao_{nome}.xlsx"),
            word_path=docx_path,
        )
        messagebox.showinfo("Sucesso", "Comparação finalizada com sucesso!")
    except Exception as e:
        messagebox.showerror("Erro", str(e))


def comparar_varios(root):
    qtd = simpledialog.askinteger("Quantidade", "Quantos pares deseja comparar?")
    if not qtd or qtd <= 0:
        return
    pasta = filedialog.askdirectory(title="Selecione a pasta de saída")
    if not pasta:
        return

    progress_window = Tk()
    progress_window.title("Progresso de Comparação")
    progress_window.geometry("400x100")

    Label(progress_window, text=f"Comparando {qtd} pares...", font=("Arial", 10)).pack(pady=10)
    progress = ttk.Progressbar(progress_window, length=350, mode="determinate", maximum=qtd)
    progress.pack(pady=5)

    for i in range(qtd):
        docx_path = filedialog.askopenfilename(
            title=f"Arquivo DOCX {i+1}", filetypes=[("Word", "*.docx")]
        )
        if not docx_path:
            continue
        url = simpledialog.askstring("URL", f"URL para o arquivo {i+1}:")
        if not url:
            continue

        try:
            docx_txt = carregar_texto_docx(docx_path)
            html_txt, main, meta, alts, titulo, imagens = carregar_texto_url(url)
            df1 = comparar_textos(docx_txt, html_txt, meta, alts)
            df1 = df1[df1["Document Text"].str.strip() != ""]
            df2 = gerar_resumo(df1)
            df3 = coletar_elementos_html(main)
            df4 = pd.DataFrame(imagens, columns=["Image URL", "Image Alt"])
            nome = re.sub(r'[\\/:*?"<>|]', '', titulo)[:50]
            salvar_em_excel(
                df1,
                df2,
                df3,
                df4,
                meta,
                url,
                os.path.join(pasta, f"comparacao_{nome}.xlsx"),
                word_path=docx_path,
            )
            progress.step(1)
            progress_window.update_idletasks()
        except Exception as e:
            messagebox.showerror("Erro", f"Erro em {docx_path}: {e}")

    messagebox.showinfo("Sucesso", "Processo de múltiplas páginas finalizado com sucesso!")
    progress_window.destroy()

if __name__ == "__main__":
    executar_comparador()